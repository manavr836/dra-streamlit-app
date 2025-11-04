#!/usr/bin/env python3
"""
DRA - Distance to Boundary v1.6

Base: v1.4 behavior preserved.
Changes in v1.6:
 - Console logging toggle (VERBOSE_LOGGING)
 - Console prints all matched candidates (for debugging)
 - CSV writes deduplicated assets per input coordinate based on Asset_Easting+Asset_Northing
 - Map plotting shows all matched assets with labels/distances; no diagonal cross lines
 - Fills a provided DOCX DRA template (Table B & Table C) per coordinate and saves DOCX and PDF
 - Configurable template path
"""


import pandas as pd
import win32com.client as win32
from datetime import date
import os


import os
import sys
import csv
import json
import logging
from typing import List, Dict, Tuple, Optional

import requests
from geopy.distance import geodesic
from pyproj import Transformer
from shapely import ops
from shapely.geometry import Point, Polygon, LineString
from shapely.ops import nearest_points

import matplotlib.pyplot as plt
try:
    import contextily as ctx
except Exception:
    ctx = None

from docx import Document  # python-docx
try:
    from docx2pdf import convert as docx2pdf_convert  # optional, Windows Word required
except Exception:
    docx2pdf_convert = None

# ====== Config ======
VERSION = "v1.6"
NO_ASSET_MESSAGE = "No asset found within 100m range"

EXPORT_MAPS = True
MAP_OUTPUT_DIR = "maps_static"
CIRCLE_RADII = [10, 25, 50, 100]   # meters (small->large)
MAP_FILE_FORMAT = "png"            # png or pdf
SEARCH_RADIUS = 100                # meters

# template settings (point to your uploaded template)
TEMPLATE_DOCX_PATH = r"C:\Users\mmogha\Downloads\DRA1_Template.docx"  # configure path to editable DOCX
EXPORT_DRA_PDF = True                       # try to convert filled docx -> pdf
CONSOLE_PRINT = True                        # override per-run (toggle printing to console)
VERBOSE_LOGGING = True                      # extra debug (prints candidate details)

# Output CSV fieldnames (single source of truth)
ZONE_A = "<10m"
ZONE_B = "10-25m"
ZONE_C = ">25-50m"
ZONE_D = ">50-100m"

OUTPUT_FIELDS = [
    "Job Reference",
    "Input Easting", "Input Northing", "Converted Lat", "Converted Lon",
    "Asset Name", "Category", "Distance(m)",
    ZONE_A, ZONE_B, ZONE_C, ZONE_D,
    "OSM_Type", "OSM_ID", "Asset_Easting", "Asset_Northing", "Raw_Tags",
    "Overpass_Server", "Exception"
]

# ====== Logging ======
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s")

# ====== Transformers ======
transformer = Transformer.from_crs("epsg:27700", "epsg:4326", always_xy=True)
transformer_back = Transformer.from_crs("epsg:4326", "epsg:27700", always_xy=True)
transformer_4326_to_3857 = Transformer.from_crs("epsg:4326", "epsg:3857", always_xy=True)
transformer_3857_to_4326 = Transformer.from_crs("epsg:3857", "epsg:4326", always_xy=True)

OVERPASS_ENDPOINTS = [
    "https://overpass.kumi.systems/api/interpreter",
    "https://overpass-api.de/api/interpreter",
    "https://lz4.overpass-api.de/api/interpreter"
]

# ====== Helpers ======
def map_category_by_tags(tags: dict, name: str = "") -> Optional[str]:
    name_l = name.lower() if name else ""
    text = " ".join([str(v).lower() for v in tags.values()]) + " " + name_l

    if tags.get("landuse") == "landfill" or tags.get("leisure") == "dump" or tags.get("waste") == "landfill":
        return "Waste Sites"
    if tags.get("landuse") == "scrap_yard" or tags.get("recycling") or "scrap" in text:
        return "Waste Sites"
    if tags.get("man_made") == "gasometer" or "gasholder" in text or "gas holder" in text:
        return "Gas holder stations"
    if tags.get("amenity") == "fuel":
        return "Petrol stations / Garages"
    if tags.get("power") == "substation" or "substation" in text:
        return "Sub-Stations"
    if tags.get("man_made") in ("wastewater_plant", "sewage_works") or "sewage" in text:
        return "Sewage Treatment Works"
    if tags.get("landuse") == "industrial" or tags.get("industrial") == "yes" or "factory" in text:
        return "Industrial / Manufacturing"
    if tags.get("landuse") == "quarry" or tags.get("mining") == "yes" or "quarry" in text:
        return "Mining (coal, metalliferous)"
    if tags.get("power") == "plant" or "plant:source" in tags or "power plant" in text:
        return "Power Plants"
    # fallbacks
    if "waste" in text:
        return "Waste Sites"
    if "petrol" in text or "garage" in text or "garages" in text or "fuel" in text:
        return "Petrol stations / Garages"
    if "substation" in text:
        return "Sub-Stations"
    if "sewage" in text or "treatment works" in text:
        return "Sewage Treatment Works"
    if "industrial" in text or "manufacturing" in text or "factory" in text:
        return "Industrial / Manufacturing"
    if "mining" in text or "quarry" in text:
        return "Mining (coal, metalliferous)"
    return None


def zone_flags(dist_m: float) -> Dict[str, str]:
    return {
        ZONE_A: "X" if dist_m < 10 else "",
        ZONE_B: "X" if 10 <= dist_m <= 25 else "",
        ZONE_C: "X" if 25 < dist_m <= 50 else "",
        ZONE_D: "X" if 50 < dist_m <= 100 else ""
    }


def query_overpass(query: str) -> Tuple[dict, str]:
    last_exc = None
    for url in OVERPASS_ENDPOINTS:
        try:
            logging.debug("Query Overpass -> %s", url)
            resp = requests.get(url, params={"data": query}, timeout=60)
            resp.raise_for_status()
            return resp.json(), url
        except Exception as e:
            logging.warning("Overpass server %s failed: %s", url, repr(e))
            last_exc = e
    raise last_exc


def query_osm(lat: float, lon: float, radius: int = SEARCH_RADIUS) -> Tuple[dict, str]:
    q = f"""
    [out:json][timeout:60];
    (
      node(around:{radius},{lat},{lon});
      way(around:{radius},{lat},{lon});
      relation(around:{radius},{lat},{lon});
    );
    out tags center geom;
    """
    return query_overpass(q)


def fetch_relation_geometry_or_member_geoms(rel_id: int) -> List[List[Tuple[float, float]]]:
    results = []
    q1 = f"[out:json][timeout:60]; relation({rel_id}); out geom;"
    try:
        resp, _ = query_overpass(q1)
        for el in resp.get("elements", []):
            if el.get("type") == "relation" and el.get("id") == rel_id and el.get("geometry"):
                coords = [(p["lon"], p["lat"]) for p in el["geometry"]]
                if coords:
                    results.append(coords)
                    return results
    except Exception:
        pass
    q2 = f"[out:json][timeout:60]; relation({rel_id}); >; out geom;"
    try:
        resp, _ = query_overpass(q2)
        for el in resp.get("elements", []):
            if el.get("type") in ("way",) and el.get("geometry"):
                coords = [(p["lon"], p["lat"]) for p in el["geometry"]]
                if coords:
                    results.append(coords)
    except Exception:
        pass
    return results


def geodesic_circle(lat: float, lon: float, radius_m: float, n_points: int = 72) -> List[Tuple[float, float]]:
    from geopy.distance import geodesic as _geodesic
    pts = []
    step = max(1, int(360 / n_points))
    for angle in range(0, 360, step):
        p = _geodesic(meters=radius_m).destination((lat, lon), angle)
        pts.append((p.longitude, p.latitude))
    return pts


def shapely_transform_geom(geom, transformer_obj: Transformer):
    return ops.transform(lambda x, y, z=None: transformer_obj.transform(x, y), geom)


# ====== Map export ======
def export_map_static(lat: float, lon: float, easting: float, northing: float, file_prefix: str,
                      assets: Optional[List[dict]] = None, no_assets: bool = False):
    os.makedirs(MAP_OUTPUT_DIR, exist_ok=True)
    fig, ax = plt.subplots(figsize=(6, 6))

    # Plot dig point
    ax.plot(lon, lat, "ro", markersize=10, label="Dig Point")

    # Draw geodesic circles
    colors = ["green", "orange", "red", "purple"]
    for r, color in zip(CIRCLE_RADII, colors):
        circ = geodesic_circle(lat, lon, r)
        if circ:
            xs, ys = zip(*circ)
            ax.plot(xs, ys, color=color, lw=2, label=f"{r}m")

    # Plot found assets
    if assets:
        for asset in assets:
            alon = asset.get("lon"); alat = asset.get("lat")
            if alon is None or alat is None:
                continue
            ax.plot(alon, alat, "bo", markersize=6, zorder=20)
            label = f"{asset.get('name','')} – {round(asset.get('dist', 0), 1)} m"
            ax.text(alon + 0.00006, alat + 0.00003, label, fontsize=7, color="blue", zorder=25)

    # Add axis labels and ticks (user requested)
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    ax.tick_params(axis='both', which='major', labelsize=8)

    # Extent: tight around 100m approx
    deg_offset = 0.0025
    ax.set_xlim(lon - deg_offset, lon + deg_offset)
    ax.set_ylim(lat - deg_offset, lat + deg_offset)

    # Basemap if available
    if ctx:
        try:
            ctx.add_basemap(ax, crs="EPSG:4326", source=ctx.providers.OpenStreetMap.Mapnik)
        except Exception as e:
            logging.debug("contextily basemap failed: %s", e)

    # Title includes coordinates
    ax.set_title(f"Dig Point with Risk Zones\n({int(easting)}_{int(northing)})")

    ax.legend(loc="upper right")
    file_path = os.path.join(MAP_OUTPUT_DIR, f"{file_prefix}.{MAP_FILE_FORMAT}")
    plt.savefig(file_path, dpi=300, bbox_inches="tight", format=MAP_FILE_FORMAT)
    plt.close(fig)
    logging.info("Map saved: %s", file_path)


# ====== DOCX Template filling (Table B & Table C) ======
def fill_dra_docx_and_export(easting: float, northing: float, matched_items: List[Dict], template_path: str,
                             out_basename: str) -> Tuple[str, Optional[str]]:
    """
    Fill Table B (screening) and Table C (risk outcome) in the provided docx template.
    Returns (docx_out_path, pdf_out_path_or_None)
    This function expects the template table layout similar to the uploaded template,
    and attempts to find a table that contains 'Table B' heading then fill rows for the site types.
    """
    out_docx = os.path.join(MAP_OUTPUT_DIR, f"{out_basename}.docx")
    out_pdf = os.path.join(MAP_OUTPUT_DIR, f"{out_basename}.pdf")

    # Prepare flags per category
    flags_by_cat = {
        "Industrial / Manufacturing": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Gas Holder Stations": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Mining": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Petrol Stations": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Sewage Treatment Works": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Sub-Stations": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Waste Sites": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Waste Sites - Scrapyard": {ZONE_A: "", ZONE_B: "", ZONE_C: ""},
        "Waste Sites - Other": {ZONE_A: "", ZONE_B: "", ZONE_C: ""}
    }

    # Map matched_items into flags_by_cat (we'll map "Waste Sites" into appropriate rows - simple approach)
    for it in matched_items:
        csv_row = it.get("csv", {})
        cat = csv_row.get("Category")
        distm = csv_row.get("Distance(m)", 9999)
        # determine zone key (for Table B we only mark <10, 10-25, >25)
        if distm < 10:
            k = ZONE_A
        elif 10 <= distm <= 25:
            k = ZONE_B
        else:
            k = ZONE_C
        # normalize category names to keys used in template
        if cat == "Industrial / Manufacturing":
            flags_by_cat["Industrial / Manufacturing"][k] = "X"
        elif cat == "Gas Holder Stations":
            flags_by_cat["Gas Holder Stations"][k] = "X"
        elif cat == "Mining":
            flags_by_cat["Mining"][k] = "X"
        elif cat == "Petrol Stations":
            flags_by_cat["Petrol Stations"][k] = "X"
        elif cat == "Sewage Treatment Works":
            flags_by_cat["Sewage Treatment Works"][k] = "X"
        elif cat == "Sub-Stations":
            flags_by_cat["Sub-Stations"][k] = "X"
        elif cat == "Waste Sites":
            # put into generic Waste Sites row
            flags_by_cat["Waste Sites"][k] = "X"
        else:
            # fallback to generic Waste Sites
            flags_by_cat["Waste Sites"][k] = "X"

    # Determine Table C risk based on rules you gave:
    # Low Risk if only assets >25m -> ZONE_C only
    # Medium if any 10–25m
    # High if any <10m
    risk = "Low"
    for it in matched_items:
        d = it.get("csv", {}).get("Distance(m)", 9999)
        if d < 10:
            risk = "High"
            break
        if 10 <= d <= 25:
            risk = "Medium"

    # Now open the template and try to fill tables
    try:
        doc = Document(template_path)
    except Exception as e:
        logging.error("Failed to open DOCX template: %s", e)
        return "", None

    # Heuristic: find the first table that contains "Table B" in surrounding paragraphs or table text
    # We'll scan paragraphs for the heading 'Table B' and use the next table.
    table_b_table = None
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    # locate index of paragraph with 'Table B' text
    table_b_idx = None
    for i, ptxt in enumerate(paragraphs):
        if "table b" in ptxt.lower():
            table_b_idx = i
            break

    if table_b_idx is not None:
        # find the table that is after that paragraph index in the document flow
        # doc.paragraphs doesn't give direct mapping to tables -> we search the nearest following table by scanning doc._body
        # fallback approach: look for table that contains the expected site type labels
        for tbl in doc.tables:
            tbl_text = " ".join(cell.text for row in tbl.rows for cell in row.cells)
            # if any of our category names are present, pick this table
            if "Industrial" in tbl_text or "Sub-Stations" in tbl_text or "Gas" in tbl_text:
                table_b_table = tbl
                break
    else:
        # fallback: attempt to find the first table with site rows by content
        for tbl in doc.tables:
            tbl_text = " ".join(cell.text for row in tbl.rows for cell in row.cells)
            if "Industrial" in tbl_text or "Sub-Stations" in tbl_text or "Gas" in tbl_text:
                table_b_table = tbl
                break

    # If we have table_b_table, attempt to fill by searching rows
    if table_b_table is not None:
        # Each row expected like: [Site Type] [No] [<10m] [10-25m] [>25m]
        # We'll match row[0] cell text for keywords and set the correct column cell text to 'X'
        for row in table_b_table.rows:
            first_cell = row.cells[0].text.strip().lower()
            # map known keywords to our flags
            if "industrial" in first_cell:
                row.cells[2].text = flags_by_cat["Industrial / Manufacturing"].get(ZONE_A, "")
                row.cells[3].text = flags_by_cat["Industrial / Manufacturing"].get(ZONE_B, "")
                row.cells[4].text = flags_by_cat["Industrial / Manufacturing"].get(ZONE_C, "")
            elif "gas" in first_cell and "holder" in first_cell or "gas holder" in first_cell or "gas holder stations" in first_cell:
                row.cells[2].text = flags_by_cat["Gas Holder Stations"].get(ZONE_A, "")
                row.cells[3].text = flags_by_cat["Gas Holder Stations"].get(ZONE_B, "")
                row.cells[4].text = flags_by_cat["Gas Holder Stations"].get(ZONE_C, "")
            elif "mining" in first_cell:
                row.cells[2].text = flags_by_cat["Mining"].get(ZONE_A, "")
                row.cells[3].text = flags_by_cat["Mining"].get(ZONE_B, "")
                row.cells[4].text = flags_by_cat["Mining"].get(ZONE_C, "")
            elif "petrol" in first_cell or "garage" in first_cell:
                row.cells[2].text = flags_by_cat["Petrol Stations"].get(ZONE_A, "")
                row.cells[3].text = flags_by_cat["Petrol Stations"].get(ZONE_B, "")
                row.cells[4].text = flags_by_cat["Petrol Stations"].get(ZONE_C, "")
            elif "sewage" in first_cell:
                row.cells[2].text = flags_by_cat["Sewage Treatment Works"].get(ZONE_A, "")
                row.cells[3].text = flags_by_cat["Sewage Treatment Works"].get(ZONE_B, "")
                row.cells[4].text = flags_by_cat["Sewage Treatment Works"].get(ZONE_C, "")
            elif "sub-station" in first_cell or "substation" in first_cell:
                row.cells[2].text = flags_by_cat["Sub-Stations"].get(ZONE_A, "")
                row.cells[3].text = flags_by_cat["Sub-Stations"].get(ZONE_B, "")
                row.cells[4].text = flags_by_cat["Sub-Stations"].get(ZONE_C, "")
            elif "landfill" in first_cell or "waste site" in first_cell:
                # try to fill a generic waste row
                row.cells[2].text = flags_by_cat["Waste Sites"].get(ZONE_A, "")
                row.cells[3].text = flags_by_cat["Waste Sites"].get(ZONE_B, "")
                row.cells[4].text = flags_by_cat["Waste Sites"].get(ZONE_C, "")
            # else skip row
    else:
        logging.warning("Could not find Table B in template to fill - template structure may differ.")

    # Fill Table C (risk outcome) - find text 'Table C' or 'Desktop risk assessment outcome' and set a checkbox cell
    # simple approach: scan tables for 'Low Risk' 'Medium Risk' 'High Risk' and set 'X' next to the chosen one
    table_c_table = None
    for tbl in doc.tables:
        tbl_text = " ".join(cell.text for row in tbl.rows for cell in row.cells)
        if "low risk" in tbl_text.lower() and "medium risk" in tbl_text.lower() and "high risk" in tbl_text.lower():
            table_c_table = tbl
            break

    if table_c_table is not None:
        # Find which column has the chosen risk and mark 'X' in that cell
        for i, row in enumerate(table_c_table.rows):
            for j, cell in enumerate(row.cells):
                if "low risk" in cell.text.lower() and risk == "Low":
                    # mark X in this cell or the adjacent cell intended for checkbox - heuristic: replace 'Low Risk' with 'X Low Risk'
                    cell.text = "X " + cell.text
                if "medium risk" in cell.text.lower() and risk == "Medium":
                    cell.text = "X " + cell.text
                if "high risk" in cell.text.lower() and risk == "High":
                    cell.text = "X " + cell.text
    else:
        logging.warning("Could not find Table C in template to fill - template structure may differ.")

    # Replace some top-level fields: Date and Location, etc. We'll set Date of Assessment and Location
    # Heuristic: Replace first occurrences of specific placeholder tokens like 'Date of Assessment:' or 'Location of work:'
    # Iterate paragraphs and replace tokens if found
    for p in doc.paragraphs:
        txt = p.text
        if "Date of Assessment" in txt and "X" in txt:
            # replace the 'X' after the label with current date placeholder; leave as 'X' for now (user asked)
            p.text = txt.replace("X", "")  # remove placeholder 'X' so user can manually set if needed
        if "Location of work" in txt and "X" in txt:
            p.text = txt.replace("X", f"{int(easting)}_{int(northing)}")

    # Save DOCX (do not overwrite template)
    doc.save(out_docx)
    logging.info("DRA docx saved: %s", out_docx)

    pdf_path = None
    if EXPORT_DRA_PDF and docx2pdf_convert is not None:
        try:
            # docx2pdf convert requires absolute paths on Windows
            docx2pdf_convert(out_docx, out_pdf)
            pdf_path = out_pdf
            logging.info("DRA PDF exported: %s", pdf_path)
        except Exception as e:
            logging.warning("docx2pdf conversion failed: %s", e)
            pdf_path = None
    else:
        if EXPORT_DRA_PDF:
            logging.warning("docx2pdf not available; DOCX saved but PDF not created.")

    return out_docx, pdf_path


# ====== CSV helper ======
def open_csv_with_sniffing(path: str):
    fh = open(path, newline='', encoding='utf-8-sig')
    sample = fh.read(4096)
    fh.seek(0)
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
    except Exception:
        dialect = csv.get_dialect("excel")
    reader = csv.DictReader(fh, dialect=dialect)
    easting_key = northing_key = None
    for k in (reader.fieldnames or []):
        if k is None:
            continue
        lk = k.strip().lower()
        if lk in ("easting", "eastings", "east"):
            easting_key = k
        if lk in ("northing", "northings", "north"):
            northing_key = k
    return fh, reader, easting_key, northing_key


# ====== Main processing ======
def process_coords(input_csv: str, output_csv: str, template_docx: Optional[str] = None, console_print: bool = True):
    fh, reader, easting_key, northing_key = open_csv_with_sniffing(input_csv)
    if not easting_key or not northing_key:
        logging.error("Couldn't find 'Easting' and 'Northing' columns in input CSV headers: %s", reader.fieldnames)
        fh.close()
        raise SystemExit("Input CSV must contain columns named 'Easting' and 'Northing' (case-insensitive).")

    os.makedirs(MAP_OUTPUT_DIR, exist_ok=True)
    with open(output_csv, 'w', newline='', encoding='utf-8-sig') as fout:
        writer = csv.DictWriter(fout, fieldnames=OUTPUT_FIELDS)
        writer.writeheader()

        rows = list(reader)
        logging.info("[%d] input rows", len(rows))

        for idx, row in enumerate(rows, start=1):
            logging.info("[%d/%d] Processing input: %s, %s", idx, len(rows), row.get(easting_key), row.get(northing_key))
            try:
                easting = float(row.get(easting_key))
                northing = float(row.get(northing_key))

            except Exception as e:
                logging.exception("Invalid Easting/Northing for row: %s", row)
                writer.writerow({
                    "Job Reference" : row.get("Job Reference"),
                    "Input Easting": row.get(easting_key), "Input Northing": row.get(northing_key),
                    "Converted Lat": "", "Converted Lon": "", "Asset Name": "", "Category": "", "Distance(m)": "",
                    ZONE_A: "", ZONE_B: "", ZONE_C: "", ZONE_D: "",
                    "OSM_Type": "", "OSM_ID": "", "Asset_Easting": "", "Asset_Northing": "", "Raw_Tags": "",
                    "Overpass_Server": "", "Exception": f"Invalid input coords: {e}"
                })
                continue
            lon, lat = transformer.transform(easting, northing)
            dig_pt_4326 = Point(lon, lat)
            dig_pt_3857 = shapely_transform_geom(dig_pt_4326, transformer_4326_to_3857)

            # Query OSM
            try:
                data, used_server = query_osm(lat, lon, radius=SEARCH_RADIUS)
                logging.info("Used Overpass server: %s", used_server)
            except Exception as e:
                logging.exception("Overpass query failed for %s,%s", lat, lon)
                writer.writerow({
                    "Job Reference": row.get("Job Reference"),
                    "Input Easting": easting, "Input Northing": northing, "Converted Lat": lat, "Converted Lon": lon,
                    "Asset Name": "", "Category": "", "Distance(m)": "",
                    ZONE_A: "", ZONE_B: "", ZONE_C: "", ZONE_D: "",
                    "OSM_Type": "", "OSM_ID": "", "Asset_Easting": "", "Asset_Northing": "", "Raw_Tags": "",
                    "Overpass_Server": "", "Exception": f"Overpass error: {e}"
                })
                if EXPORT_MAPS:
                    export_map_static(lat, lon, easting, northing, f"{row.get('Job Reference')}_error", assets=None, no_assets=True)
                continue

            elements = data.get("elements", []) or []
            logging.info("Found %d candidate elements within %dm", len(elements), SEARCH_RADIUS)

            matched_items_all: List[Dict] = []  # all matched candidates (console)
            matched_items_unique: List[Dict] = []  # unique by asset_easting/northing (CSV)

            for el in elements:
                tags = el.get("tags", {}) or {}
                name = tags.get("name", "Unnamed asset")
                category = map_category_by_tags(tags, name)
                if not category:
                    continue

                el_type = el.get("type")
                el_id = el.get("id")

                dist_m = None
                nearest_lon = None
                nearest_lat = None

                # node
                if el_type == "node" and "lat" in el and "lon" in el:
                    el_lat = el["lat"]; el_lon = el["lon"]
                    dist_m = geodesic((lat, lon), (el_lat, el_lon)).meters
                    nearest_lat, nearest_lon = el_lat, el_lon

                # way
                elif el_type == "way" and "geometry" in el:
                    coords = [(p["lon"], p["lat"]) for p in el["geometry"]]
                    if coords:
                        try:
                            if len(coords) >= 3 and coords[0] == coords[-1]:
                                geom_4326 = Polygon(coords); target_4326 = geom_4326.exterior
                            else:
                                geom_4326 = LineString(coords); target_4326 = geom_4326
                        except Exception:
                            continue
                        try:
                            geom_3857 = shapely_transform_geom(target_4326, transformer_4326_to_3857)
                            _, nearest_on_geom = nearest_points(dig_pt_3857, geom_3857)
                            nx, ny = nearest_on_geom.x, nearest_on_geom.y
                            nearest_lon, nearest_lat = transformer_3857_to_4326.transform(nx, ny)
                            dist_m = geodesic((lat, lon), (nearest_lat, nearest_lon)).meters
                        except Exception:
                            continue

                # relation
                elif el_type == "relation":
                    # try geometry inline
                    if "geometry" in el and el.get("geometry"):
                        coords = [(p["lon"], p["lat"]) for p in el["geometry"]]
                        if coords:
                            try:
                                if len(coords) >= 3 and coords[0] == coords[-1]:
                                    geom_4326 = Polygon(coords); target_4326 = geom_4326.exterior
                                else:
                                    geom_4326 = LineString(coords); target_4326 = geom_4326
                            except Exception:
                                target_4326 = None
                            if target_4326 is not None:
                                try:
                                    geom_3857 = shapely_transform_geom(target_4326, transformer_4326_to_3857)
                                    _, nearest_on_geom = nearest_points(dig_pt_3857, geom_3857)
                                    nx, ny = nearest_on_geom.x, nearest_on_geom.y
                                    nearest_lon, nearest_lat = transformer_3857_to_4326.transform(nx, ny)
                                    dist_m = geodesic((lat, lon), (nearest_lat, nearest_lon)).meters
                                except Exception:
                                    pass

                    # if still not found, fetch member geometries and pick nearest
                    if dist_m is None:
                        found_geom_lists = fetch_relation_geometry_or_member_geoms(el_id)
                        best_dist = None
                        best_nearest = (None, None)
                        for coords in found_geom_lists:
                            if not coords:
                                continue
                            try:
                                if len(coords) >= 3 and coords[0] == coords[-1]:
                                    g4326 = Polygon(coords); target_4326 = g4326.exterior
                                else:
                                    g4326 = LineString(coords); target_4326 = g4326
                            except Exception:
                                continue
                            try:
                                g3857 = shapely_transform_geom(target_4326, transformer_4326_to_3857)
                                _, nearest_on_g = nearest_points(dig_pt_3857, g3857)
                                nx, ny = nearest_on_g.x, nearest_on_g.y
                                nlon, nlat = transformer_3857_to_4326.transform(nx, ny)
                                d = geodesic((lat, lon), (nlat, nlon)).meters
                                if best_dist is None or d < best_dist:
                                    best_dist = d
                                    best_nearest = (nlon, nlat)
                            except Exception:
                                continue
                        if best_dist is not None:
                            dist_m = best_dist
                            nearest_lon, nearest_lat = best_nearest

                # fallback center
                if dist_m is None and el.get("center"):
                    c = el["center"]
                    if c.get("lat") is not None and c.get("lon") is not None:
                        nearest_lat = c.get("lat"); nearest_lon = c.get("lon")
                        dist_m = geodesic((lat, lon), (nearest_lat, nearest_lon)).meters

                if dist_m is None:
                    continue

                if dist_m > SEARCH_RADIUS:
                    continue

                flags = zone_flags(dist_m)
                if not any(flags.values()):
                    continue

                try:
                    asset_easting, asset_northing = transformer_back.transform(nearest_lon, nearest_lat)
                except Exception:
                    asset_easting, asset_northing = None, None

                dist_round = round(dist_m, 2)

                out_row = {
                    "Job Reference": row.get("Job Reference"),
                    "Input Easting": easting,
                    "Input Northing": northing,
                    "Converted Lat": lat,
                    "Converted Lon": lon,
                    "Asset Name": name,
                    "Category": category,
                    "Distance(m)": dist_round,
                    ZONE_A: flags[ZONE_A],
                    ZONE_B: flags[ZONE_B],
                    ZONE_C: flags[ZONE_C],
                    ZONE_D: flags[ZONE_D],
                    "OSM_Type": el_type,
                    "OSM_ID": el_id,
                    "Asset_Easting": round(asset_easting, 2) if asset_easting is not None else "",
                    "Asset_Northing": round(asset_northing, 2) if asset_northing is not None else "",
                    "Raw_Tags": json.dumps(tags, ensure_ascii=False),
                    "Overpass_Server": used_server,
                    "Exception": ""
                }

                # append to matched_items_all for console
                matched_items_all.append({
                    "csv": out_row,
                    "plot_lon": nearest_lon,
                    "plot_lat": nearest_lat,
                    "plot_name": name,
                    "plot_dist": dist_round
                })

                # console debug
                if console_print or VERBOSE_LOGGING:
                    logging.info("Candidate processed: id=%s type=%s name=%s dist=%.2fm flags=%s tags=%s",
                                 el_id, el_type, name, dist_round, flags, tags)

            # dedupe for CSV: unique by Asset_Easting & Asset_Northing (string compare)
            seen_coords = set()
            for item in matched_items_all:
                csv_row = item["csv"]
                ae = csv_row.get("Asset_Easting", "")
                an = csv_row.get("Asset_Northing", "")
                key = f"{ae}_{an}"
                if ae == "" or an == "":
                    # keep assets without easting/northing (edge case) - add with unique id
                    key = key + "_" + str(csv_row.get("OSM_ID", ""))
                if key in seen_coords:
                    # skip duplicate for CSV output, but console already printed all
                    continue
                seen_coords.add(key)
                matched_items_unique.append(item)

            # Write CSV rows (from matched_items_unique)
            if matched_items_unique:
                logging.info("Selected %d unique assets for output", len(matched_items_unique))
                for it in matched_items_unique:
                    safe_row = {k: it["csv"].get(k, "") for k in OUTPUT_FIELDS}
                    writer.writerow(safe_row)
            else:
                logging.info("No assets found within %dm for this dig point", SEARCH_RADIUS)
                writer.writerow({
                    "Job Reference": row.get("Job Reference"),
                    "Input Easting": easting, "Input Northing": northing, "Converted Lat": lat, "Converted Lon": lon,
                    "Asset Name": "", "Category": "", "Distance(m)": "",
                    ZONE_A: "", ZONE_B: "", ZONE_C: "", ZONE_D: "",
                    "OSM_Type": "", "OSM_ID": "", "Asset_Easting": "", "Asset_Northing": "", "Raw_Tags": "",
                    "Overpass_Server": used_server if 'used_server' in locals() else "",
                    "Exception": NO_ASSET_MESSAGE
                })

            # Map plotting: include all unique items (plot points) so map not overcrowded with exact duplicates
            assets_for_map = []
            for it in matched_items_unique:
                if it.get("plot_lon") is None or it.get("plot_lat") is None:
                    continue
                assets_for_map.append({
                    "lon": it.get("plot_lon"),
                    "lat": it.get("plot_lat"),
                    "name": it.get("plot_name", ""),
                    "dist": it.get("plot_dist", 0)
                })

            if EXPORT_MAPS:
                export_map_static(
                    lat,
                    lon,
                    easting,
                    northing,
                    f"{row.get('Job Reference')}",
                    assets=assets_for_map if assets_for_map else None,
                    no_assets=(len(matched_items_all) == 0)
                )

            # Create DRA docs (Table B, Table C) if template provided
            if template_docx:
                pdf_name = f"DRA_{int(easting)}_{int(northing)}"
                try:
                    out_docx, out_pdf = fill_dra_docx_and_export(easting, northing, matched_items_all, template_docx, pdf_name)
                    logging.info("Table B/C DRA documents created: %s , %s", out_docx, out_pdf)
                except Exception as e:
                    logging.error("Failed to create Table B/C documents: %s", e)

        logging.info("Processing finished. CSV saved to %s", output_csv)
    # close input file handle
    fh.close()
    logging.info("=== DRA Processing finished for all coordinates ===")

############################################################################################## Creating the WordFile ########################################################################################

def CreatingWordFile(excel_path,template_path,output_dir):
    # ------------------- CONFIG -------------------


    # Mapping Category keywords in Excel to site names in Word
    category_map = {
        "Industrial / Manufacturing": "Industrial",
        "Gas holder stations": "Gas",
        "Mining (coal, metalliferous)": "Mining",
        "Petrol stations / Garages": "Petrol",
        "Sewage Treatment Works": "Sewage",
        "Sub-Stations": "Sub_Stat",
        "Waste Site – Landfill & Treatment / Disposal": "Landfill",
        "Waste Site – Scrapyard / Metal Recycling": "Scrapyard",
        "Waste Sites": "Other",  # "Waste Site – Other": "Other",
    }

    # Map answer to checkbox suffix
    answer_suffix_map = {
        "No": "NO",
        "<10m": "10",
        "10-25m": "10_25",
        ">25m": "_25",
    }
    # ------------------------------------------------

    # Load Excel
    df = pd.read_csv(excel_path)

    # Group by coordinates
    grouped = df.groupby(["Input Easting", "Input Northing"])

    # Start Word once
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False

    for (easting, northing), group in grouped:
        print(f"\n📍 Processing {easting}, {northing}")

        # Open fresh template for each coordinate set
        doc = word.Documents.Open(template_path)

        # Build site_answers dynamically
        site_answers = {site: "No" for site in category_map.values()}  # default all to No

        # ✅ Aggregate to get minimum distance per category
        category_min_dist = (
            group.groupby("Category")["Distance(m)"]
            .min()
            .to_dict()
        )

        # Process each unique category using its min distance
        for category, distance in category_min_dist.items():
            if pd.isna(category):
                continue  # stays as "No"

            site_key = category_map.get(category.strip())
            if not site_key:
                continue

            # Decide distance band
            if site_key == "Other":
                if distance <= 10:
                    site_answers[site_key] = "<10m"
                    site_answers["Landfill"] = "<10m"
                    site_answers["Scrapyard"] = "<10m"
                elif distance <= 25:
                    site_answers[site_key] = "10-25m"
                    site_answers["Landfill"] = "10-25m"
                    site_answers["Scrapyard"] = "10-25m"
                elif distance > 25:
                    site_answers[site_key] = ">25m"
                    site_answers["Landfill"] = ">25m"
                    site_answers["Scrapyard"] = ">25m"
            else:
                if distance <= 10:
                    site_answers[site_key] = "<10m"
                elif distance <= 25:
                    site_answers[site_key] = "10-25m"
                elif distance > 25:
                    site_answers[site_key] = ">25m"

            print(f"   ➡️ {site_key}: {site_answers[site_key]} (min distance = {distance})")

        # Tick Table B checkboxes (Content Controls)
        for cc in doc.ContentControls:
            if cc.Type == 8:  # Checkbox
                for site, answer in site_answers.items():
                    expected_suffix = answer_suffix_map.get(answer)
                    expected_tag = f"{site}{expected_suffix}"
                    if cc.Tag == expected_tag:
                        cc.Checked = True
                    else:
                        # keep all others unchecked
                        if cc.Tag.startswith(site):
                            cc.Checked = False

        # ------------------- RISK LOGIC -------------------
        answers = site_answers.values()
        if any(ans == "<10m" for ans in answers):
            risk_level = "HighRisk"
        elif any(ans == "10-25m" for ans in answers):
            risk_level = "MediumRisk"
        else:
            risk_level = "LowRisk"

        print(f"   📊 Risk determined: {risk_level}")

        # Tick risk checkbox (Content Controls)
        for cc in doc.ContentControls:
            if cc.Type == 8 and cc.Tag in ["HighRisk", "MediumRisk", "LowRisk"]:
                cc.Checked = (cc.Tag == risk_level)

        # ------------------- Update Date -------------------
        for tbl in doc.Tables:
            for cell in tbl.Range.Cells:
                if "date of assessment" in cell.Range.Text.lower():
                    next_col = cell.ColumnIndex + 1
                    for c in tbl.Range.Cells:
                        if c.RowIndex == cell.RowIndex and c.ColumnIndex == next_col:
                            c.Range.Text = date.today().strftime("%d-%m-%Y")
                            break

        # ------------------- Save -------------------
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"DRA1_{group['Job Reference'].iloc[0]}.docx")
        doc.SaveAs(output_path)
        doc.Close()

        print(f"   💾 Saved: {output_path}")

    word.Quit()
    print("\n✅ Processing complete.")


######################################################################### Creating Word File Ended ######################################################################

# ===== Entry point =====
if __name__ == "__main__":
    logging.info("DRA Distance to Boundary %s", VERSION)
    logging.info("Python version: %s", sys.version.replace("\n", " "))
    # call with template path and console toggle - change as needed
    template = TEMPLATE_DOCX_PATH if os.path.exists(TEMPLATE_DOCX_PATH) else None
    if not template:
        logging.warning("Template DOCX path not found or not provided - Table B/C will not be produced.")

    excel_path = r"C:\Users\mmogha\OneDrive - Sopra Steria\Desktop\SafeDig DRA\Input\input 1.csv"
    output_dir = r"C:\Users\mmogha\OneDrive - Sopra Steria\Desktop\SafeDig DRA\Output_ExcelFiles\outputnew.csv"
    template_path = r"C:\Users\mmogha\OneDrive - Sopra Steria\Desktop\DRA1_Template.docx"
    output_doc = r"C:\Users\mmogha\OneDrive - Sopra Steria\Desktop\SafeDig DRA\Output_WordFiles"

    process_coords(excel_path, output_dir, template_docx=template, console_print=CONSOLE_PRINT)

    CreatingWordFile(output_dir, template_path, output_doc)

    if EXPORT_MAPS:
        logging.info("Maps and DRA docs exported to: %s", MAP_OUTPUT_DIR)
