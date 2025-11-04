import os
import csv
from docx import Document
import pandas as pd

# Dummy placeholder for map directory (Streamlit sets this dynamically)
MAP_OUTPUT_DIR = "Maps"
TEMPLATE_DOCX_PATH = "template.docx"  # optional default path


def process_coords(input_csv, output_csv, template_docx=None, console_print=False):
    """
    Example CSV processing placeholder.
    Modify this logic to your actual coordinate processing.
    """
    df = pd.read_csv(input_csv)
    df["Processed"] = True
    df.to_csv(output_csv, index=False)
    if console_print:
        print(f"Processed CSV saved to {output_csv}")


def CreatingWordFile(output_csv, template_path, output_dir):
    """
    Generate Word files from a DOCX template and CSV input.
    Works both on Windows and Streamlit Cloud (Linux).
    """
    os.makedirs(output_dir, exist_ok=True)

    with open(output_csv, newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)

        for i, row in enumerate(reader, start=1):
            doc = Document(template_path)

            # Replace placeholders in paragraphs
            for p in doc.paragraphs:
                for key, val in row.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in p.text:
                        for run in p.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(val))

            # Replace placeholders in tables
            for table in doc.tables:
                for r in table.rows:
                    for cell in r.cells:
                        for key, val in row.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(val))

            # Save one file per record
            asset_id = row.get("AssetID", f"record_{i}")
            output_file = os.path.join(output_dir, f"{asset_id}.docx")
            doc.save(output_file)
            print(f"✅ Created: {output_file}")
